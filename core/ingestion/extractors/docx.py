"""DOCX content extractor using python-docx."""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Union

from docx import Document as DocxDocument

from core.exceptions import ExtractionError
from core.ingestion.extractors.base import BaseExtractor, ExtractedDocument

logger = logging.getLogger(__name__)


class DocxExtractor(BaseExtractor):
    """Extracts text content from Microsoft Word (.docx) files."""

    supported_extensions = [".docx"]

    async def extract(self, source: Union[str, Path, bytes]) -> ExtractedDocument:
        """Extract text from a DOCX file or byte stream.

        Args:
            source: File path or raw DOCX bytes.

        Returns:
            ExtractedDocument with paragraph text joined by newlines.

        Raises:
            ExtractionError: If the document cannot be read or parsed.
        """
        try:
            if isinstance(source, bytes):
                doc = DocxDocument(io.BytesIO(source))
                file_name = "uploaded.docx"
            else:
                path = Path(source)
                if not path.exists():
                    raise ExtractionError(f"DOCX file not found: {path}")
                doc = DocxDocument(str(path))
                file_name = path.name

            paragraphs: list[str] = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)

            content = "\n\n".join(paragraphs)

            metadata = {
                "source_type": "docx",
                "file_name": file_name,
                "paragraph_count": len(doc.paragraphs),
            }

            core_props = doc.core_properties
            title = core_props.title or file_name
            if core_props.author:
                metadata["author"] = core_props.author

            logger.info("Extracted %d paragraphs from DOCX %s", len(paragraphs), file_name)
            return ExtractedDocument(content=content, metadata=metadata, title=title)

        except ExtractionError:
            raise
        except Exception as exc:
            raise ExtractionError(f"Failed to extract DOCX: {exc}") from exc
